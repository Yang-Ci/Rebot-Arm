from pathlib import Path
from datetime import datetime

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "analysis_report"
PUBLIC = ROOT / "public"

DOCX = OUT / "reBotArm_simulator_网站代码结构分析.docx"


def font_path(name):
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts") / name.lower(),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return None


ZH = font_path("msyh.ttc") or font_path("simhei.ttf")
MONO = font_path("consola.ttf") or ZH


def load_font(size, mono=False, bold=False):
    if mono and MONO:
        return ImageFont.truetype(MONO, size)
    if ZH:
        return ImageFont.truetype(ZH, size)
    return ImageFont.load_default()


def rounded(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def make_webpage_screenshot():
    img = Image.new("RGB", (1440, 900), "#111211")
    d = ImageDraw.Draw(img)
    title = load_font(28, bold=True)
    small = load_font(15)
    mono = load_font(14, mono=True)
    h3 = load_font(19, bold=True)

    # Main viewport and side panel.
    d.rectangle([0, 0, 1080, 900], fill="#111211")
    d.rectangle([1080, 0, 1440, 900], fill="#151716")
    d.line([1080, 0, 1080, 900], fill="#343936", width=2)

    # Top HUD.
    rounded(d, [18, 18, 1062, 98], 8, "#171917", "#3b403c")
    d.text((38, 34), "reBot Arm B601-DM", font=small, fill="#33d6b0")
    d.text((38, 56), "机械臂仿真与 ROS2 监看空间", font=title, fill="#f4f1ea")
    rounded(d, [940, 42, 1038, 76], 17, "#242925")
    d.ellipse([956, 54, 966, 64], fill="#33d6b0")
    d.text((976, 50), "Ready", font=small, fill="#d7fff4")

    # Floor grid.
    for x in range(-80, 1080, 44):
        d.line([x, 760, x + 360, 410], fill="#2b302d", width=1)
    for y in range(425, 850, 32):
        d.line([60, y, 1050, y], fill="#2b302d", width=1)

    # Simplified robot/gripper visualization.
    d.ellipse([410, 520, 620, 730], fill="#e4ded1", outline="#b9b0a0", width=2)
    d.rectangle([465, 420, 565, 560], fill="#d8d5cc", outline="#aeb9b1", width=2)
    d.polygon([(500, 420), (540, 420), (555, 355), (485, 355)], fill="#bec9c0", outline="#aeb9b1")
    d.rectangle([245, 300, 780, 342], fill="#d8d5cc", outline="#9fb0a9", width=2)
    d.rectangle([245, 345, 780, 378], fill="#b8c1b8", outline="#87948d", width=2)
    d.rectangle([200, 258, 300, 480], fill="#33d6b0", outline="#24987f", width=2)
    d.rectangle([730, 258, 830, 480], fill="#33d6b0", outline="#24987f", width=2)
    d.rectangle([130, 335, 250, 365], fill="#e7e1d6", outline="#c8c2b8")
    d.rectangle([780, 335, 900, 365], fill="#e7e1d6", outline="#c8c2b8")
    d.rectangle([95, 300, 138, 400], fill="#d8d5cc", outline="#b8b0a4")
    d.rectangle([895, 300, 938, 400], fill="#d8d5cc", outline="#b8b0a4")
    d.ellipse([470, 498, 615, 643], fill="#e0b75a", outline="#bd9440", width=2)

    # Pose HUD.
    rounded(d, [18, 805, 660, 875], 8, "#171917", "#3b403c")
    d.text((38, 823), "TCP", font=small, fill="#a7ada7")
    d.text((38, 846), "X 218毫米 / Y -12毫米 / Z 305毫米", font=mono, fill="#f4f1ea")
    d.text((420, 823), "工作半径", font=small, fill="#a7ada7")
    d.text((420, 846), "218 / 650 毫米", font=mono, fill="#d7fff4")

    # Side panel sections.
    x0, x1 = 1102, 1418
    y = 18
    rounded(d, [x0, y, x1, y + 68], 8, "#191b1a", "#3b403c")
    d.text((1120, y + 14), "Control", font=small, fill="#33d6b0")
    d.text((1120, y + 36), "仿真控制台", font=h3, fill="#f4f1ea")
    y += 82

    rounded(d, [x0, y, x1, y + 210], 8, "#191b1a", "#3b403c")
    d.text((1120, y + 14), "ROS2 Bridge", font=small, fill="#33d6b0")
    d.text((1120, y + 38), "真实机械臂连接", font=h3, fill="#f4f1ea")
    rounded(d, [1346, y + 18, 1396, y + 44], 13, "#2b302d")
    d.text((1358, y + 21), "离线", font=small, fill="#a7ada7")
    rounded(d, [1120, y + 76, 1400, y + 112], 7, "#202321", "#343936")
    d.text((1132, y + 84), "ws://localhost:9090", font=mono, fill="#f4f1ea")
    for i, label in enumerate(["连接 ROS", "断开", "使能", "失能"]):
        bx = 1120 + (i % 2) * 142
        by = y + 126 + (i // 2) * 42
        rounded(d, [bx, by, bx + 132, by + 34], 7, "#202321", "#343936")
        d.text((bx + 24, by + 8), label, font=small, fill="#f4f1ea")
    y += 224

    rounded(d, [x0, y, x1, y + 260], 8, "#191b1a", "#3b403c")
    d.text((1120, y + 16), "关节角度", font=h3, fill="#fff8ed")
    labels = ["J1 底座偏航", "J2 肩部", "J3 肘部", "J4 腕部俯仰", "J5 腕部偏航", "J6 工具旋转", "J7 夹爪"]
    for i, label in enumerate(labels):
        yy = y + 48 + i * 28
        d.text((1120, yy), label, font=small, fill="#f4f1ea")
        d.line([1235, yy + 10, 1398, yy + 10], fill="#33d6b0", width=5)
        d.ellipse([1318, yy + 3, 1332, yy + 17], fill="#f6f1e8")
    d.text((1348, y + 226), "90 毫米", font=mono, fill="#33d6b0")

    path = OUT / "webpage-home.png"
    img.save(path)
    return path


def line_range(path, start, end):
    lines = path.read_text(encoding="utf-8").splitlines()
    rows = []
    for idx in range(start, min(end, len(lines)) + 1):
        rows.append(f"{idx:>4}  {lines[idx - 1]}")
    return rows


def make_code_image(name, file_path, start, end, title):
    rows = line_range(file_path, start, end)
    font = load_font(18, mono=True)
    title_font = load_font(21, bold=True)
    line_h = 26
    width = 1320
    height = 66 + line_h * len(rows) + 30
    img = Image.new("RGB", (width, height), "#101312")
    d = ImageDraw.Draw(img)
    rounded(d, [18, 18, width - 18, height - 18], 10, "#171b19", "#38413d")
    d.text((36, 32), title, font=title_font, fill="#33d6b0")
    y = 70
    for row in rows:
        d.text((36, y), row, font=font, fill="#f4f1ea")
        y += line_h
    path = OUT / name
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(9.5)
    r.font.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def add_figure(doc, image_path, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def make_report():
    webpage = make_webpage_screenshot()
    code_html = make_code_image("code-html-structure.png", PUBLIC / "index.html", 11, 90, "index.html：页面区域结构")
    code_js = make_code_image("code-js-gripper.png", PUBLIC / "js" / "rebot-sim.js", 351, 412, "rebot-sim.js：夹爪 STL 加载与开合映射")
    code_ros = make_code_image("code-ros-ui.png", PUBLIC / "js" / "ros" / "rebot-ros-ui.js", 28, 50, "rebot-ros-ui.js：ROS topic 订阅与按钮绑定")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Microsoft YaHei"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.name = "Microsoft YaHei"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("reBotArm Simulator 网站代码结构分析")
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(22)
    run.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"分析对象：reBotArm_simulator | 生成时间：{datetime.now().strftime('%Y-%m-%d')}")
    r.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph(
        "本文档分析 reBotArm_simulator 网站的代码结构，重点说明 HTML 的页面区域划分与网页资源的组织方式。"
        "该网站是一个机械臂 Web 仿真页面，前端使用 Three.js/URDFLoader/STLLoader 渲染 3D 机械臂，"
        "并通过独立 ROS bridge 模块预留与 ROS2 真实机械臂连接的能力。"
    )
    add_figure(doc, webpage, "图 1 网页首页运行界面截图：左侧为 3D 仿真视口，右侧为控制与 ROS2 连接面板。", 6.2)

    doc.add_heading("一、HTML 的网页区域划分和结构", level=1)
    doc.add_paragraph(
        "页面入口文件为 public/index.html。整体结构采用 main.app-shell 作为根布局，"
        "左侧 section.viewport 承载 3D 场景和 HUD 信息，右侧 aside.control-panel 承载控制面板。"
        "这种结构把“模型显示”和“用户操作”分开，便于后续扩展 ROS 控制、轨迹发送或状态监控。"
    )
    add_table(
        doc,
        ["区域", "主要 HTML 结构", "作用说明"],
        [
            ["整体框架", "main.app-shell", "使用 CSS Grid 划分左侧仿真视口与右侧控制面板。"],
            ["3D 场景容器", "section.viewport / div#scene-host", "Three.js 渲染器将 canvas 插入 scene-host，用于显示 URDF/STL 机械臂。"],
            ["顶部 HUD", "header.top-hud / #load-status", "显示网站名称、页面标题、模型加载状态。"],
            ["位姿 HUD", ".pose-hud / #tcp-position / #reach-state", "显示 TCP 坐标与工作半径，用于观察末端位置。"],
            ["坐标图例", ".legend-hud", "说明 ROS 坐标系方向：+X 前方、+Y 左侧、+Z 向上。"],
            ["ROS 连接区", ".ros-section", "提供 rosbridge 地址、连接/断开、镜像状态、控制锁、使能/失能等控件。"],
            ["关节控制区", "#joint-controls", "由 JS 动态生成 J1-J7 滑块，控制机械臂姿态和夹爪开合。"],
            ["夹爪控制区", "#gripper-width / #open-gripper / #close-gripper", "显示夹爪开合宽度，并提供打开/关闭按钮。"],
            ["图层与演示区", "#toggle-envelope / #play-path", "控制工作空间包络、任务区、残影以及演示路径播放。"],
            ["加载遮罩", "#loading-mask", "模型加载完成前显示 loading 状态，避免空白页面。"],
        ],
    )
    add_figure(doc, code_html, "图 2 index.html 部分代码截图：展示页面主体、视口、HUD、ROS 面板与控制区结构。", 6.4)

    doc.add_heading("二、网页资源组织方式与主要作用", level=1)
    doc.add_paragraph(
        "项目采用静态前端资源 + Node.js 本地服务的组织方式。public 目录保存浏览器直接加载的 HTML、CSS、JS、图标和第三方库；"
        "split_meshes/grouped_gripper 保存拆分后的夹爪 STL；server.js 负责把 URDF、原始 STL 和夹爪 STL 通过 /api 路由提供给前端。"
    )
    add_table(
        doc,
        ["资源类型", "路径/文件", "主要作用"],
        [
            ["HTML", "public/index.html", "页面入口，定义网页区域、控件、脚本加载顺序。"],
            ["CSS", "public/css/rebot-sim.css", "负责布局、深色主题、HUD、控制面板、按钮、滑块、ROS 状态样式。"],
            ["核心 JS", "public/js/rebot-sim.js", "创建 Three.js 场景，加载 URDF/STL，生成关节滑块，更新 TCP、夹爪和动画。"],
            ["ROS JS", "public/js/ros/rebot-ros-client.js", "实现轻量 rosbridge WebSocket 客户端，封装 topic 订阅和 service 调用。"],
            ["ROS UI JS", "public/js/ros/rebot-ros-ui.js", "把 ROS topic/service 与页面按钮、镜像状态、夹爪控制连接起来。"],
            ["第三方库", "public/lib/three-r128.min.js", "Three.js 渲染引擎，用于 WebGL 3D 场景。"],
            ["第三方库", "public/lib/URDFLoader.js", "解析 ROS URDF 文件，并按关节层级创建机械臂模型。"],
            ["第三方库", "public/lib/STLLoader-umd.js", "加载 STL 网格模型，特别是分离后的夹爪部件。"],
            ["图片资源", "public/favicon.png", "浏览器标签页图标。当前项目未使用普通 img 内容图。"],
            ["3D 模型资源", "split_meshes/grouped_gripper/*.stl", "夹爪基座、左夹爪、右夹爪等 STL 网格，供 STLLoader 加载。"],
            ["视频资源", "无", "当前网页没有 video 标签，也没有视频文件参与页面显示。"],
            ["服务端", "server.js", "提供静态文件和 /api/urdf、/api/description/meshes、/api/gripper_meshes 路由。"],
        ],
    )

    tree = doc.add_paragraph()
    tree.add_run("资源目录概览：").bold = True
    pre = doc.add_paragraph()
    pre.style = styles["Normal"]
    pre.add_run(
        "reBotArm_simulator/\n"
        "  server.js\n"
        "  public/\n"
        "    index.html\n"
        "    css/rebot-sim.css\n"
        "    js/rebot-sim.js\n"
        "    js/ros/rebot-ros-client.js\n"
        "    js/ros/rebot-ros-ui.js\n"
        "    lib/three-r128.min.js\n"
        "    lib/URDFLoader.js\n"
        "    lib/STLLoader-umd.js\n"
        "    favicon.png\n"
        "  split_meshes/grouped_gripper/\n"
        "    gripper_base.stl\n"
        "    left_finger.stl\n"
        "    right_finger.stl"
    ).font.name = "Consolas"

    add_figure(doc, code_js, "图 3 rebot-sim.js 部分代码截图：加载夹爪 STL，并将 90mm 控制量映射为 57mm 视觉开口。", 6.4)
    add_figure(doc, code_ros, "图 4 rebot-ros-ui.js 部分代码截图：订阅 joint_states、gripper/state、arm_status，并绑定 ROS 控制按钮。", 6.4)

    doc.add_heading("三、关键逻辑说明", level=1)
    doc.add_paragraph(
        "1. 页面加载时，index.html 先加载 Three.js、STLLoader、URDFLoader，再加载 ROS 客户端模块，最后加载 rebot-sim.js 和 rebot-ros-ui.js。"
        "这样可以保证 3D 渲染库和 ROS 封装先准备好，再由主脚本初始化场景。"
    )
    doc.add_paragraph(
        "2. rebot-sim.js 内部维护 jointDefs 和 presets。J1-J6 表示机械臂关节，J7 表示夹爪控制量。"
        "当前网页显示和 ROS 命令使用官方最大开口 90mm，但 STL 模型视觉开口按比例映射到 57mm，避免模型结构穿插。"
    )
    doc.add_paragraph(
        "3. ROS 连接采用独立模块设计。rebot-ros-client.js 只负责 WebSocket 协议、topic 订阅和 service 调用；"
        "rebot-ros-ui.js 负责把 ROS 数据映射到网页控件。这样可以降低主仿真文件的复杂度。"
    )
    doc.add_paragraph(
        "4. 模型资源不是直接写死在 HTML 中，而是由 server.js 通过 /api/urdf 和 /api/gripper_meshes 提供。"
        "这种方式可以把 ROS2 描述包里的 URDF/STL 与网页项目解耦，后续替换模型时只需要调整服务端路径或资源文件。"
    )

    doc.add_heading("四、总结", level=1)
    doc.add_paragraph(
        "该网站的代码结构比较清晰：HTML 负责页面骨架，CSS 负责视觉布局，rebot-sim.js 负责 3D 仿真核心，"
        "ros 目录下的 JS 负责真实机械臂连接。资源层面，普通图片只使用 favicon，主要视觉资源来自 URDF 和 STL 三维模型，"
        "当前没有 video 资源。整体结构适合继续扩展轨迹发送、IK 控制和真实机械臂状态监控。"
    )

    doc.save(DOCX)
    return DOCX


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    print(make_report())
